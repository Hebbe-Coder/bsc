"""
Document Parser - 文档解析器

支持格式：
    - .docx (Word文档)
    - .pdf (PDF文档)
    - .txt (纯文本)
    - .png/.jpg/.jpeg/.gif/.webp (图片文件，通过OCR识别)

功能：
    1. 根据文件扩展名自动选择解析器
    2. 提取文档文本内容
    3. 返回带来源标签的结构化文本
    4. 支持多文件合并
    5. 支持OCR处理扫描件和图片
"""
from __future__ import annotations
import os, logging, traceback, base64, threading

logger = logging.getLogger(__name__)

from app.core.config import settings

SUPPORTED_EXTENSIONS = [".docx", ".pdf", ".txt", ".png", ".jpg", ".jpeg", ".gif", ".webp"]
MAX_FILE_SIZE = settings.MAX_FILE_SIZE_MB * 1024 * 1024


class DocumentParser:
    """文档解析器"""

    def __init__(self, llm_service=None):
        self._docx_available = False
        self._pdf_available = False
        self._ocr_available = False
        self._llm_service = llm_service
        self._try_import_dependencies()

    def _try_import_dependencies(self):
        """尝试导入依赖"""
        global docx, pdfplumber
        try:
            import docx
            self._docx_available = True
        except ImportError:
            logger.warning("python-docx not installed, .docx parsing will be disabled")

        try:
            import pdfplumber
            self._pdf_available = True
        except ImportError:
            logger.warning("pdfplumber not installed, .pdf parsing will be disabled")

        try:
            import fitz
            self._fitz_available = True
        except ImportError:
            logger.warning("pymupdf not installed, PDF OCR fallback will be disabled")

        try:
            from PIL import Image
            self._pil_available = True
        except ImportError:
            logger.warning("Pillow not installed, image processing will be disabled")

    def set_llm_service(self, llm_service):
        """设置LLM服务用于OCR"""
        self._llm_service = llm_service

    def parse_file(self, file_bytes: bytes, filename: str) -> dict:
        """
        解析单个文件
        
        Args:
            file_bytes: 文件字节内容
            filename: 文件名（用于判断格式）
        
        Returns:
            dict: {"success": bool, "text": str, "filename": str, "error": str}
        """
        ext = os.path.splitext(filename)[1].lower()

        if ext not in SUPPORTED_EXTENSIONS:
            return {
                "success": False,
                "text": "",
                "filename": filename,
                "error": f"不支持的文件格式: {ext}，仅支持: {', '.join(SUPPORTED_EXTENSIONS)}",
            }

        if len(file_bytes) > MAX_FILE_SIZE:
            return {
                "success": False,
                "text": "",
                "filename": filename,
                "error": f"文件大小超过限制({MAX_FILE_SIZE//1024//1024}MB)",
            }

        try:
            if ext == ".docx":
                return self._parse_docx(file_bytes, filename)
            elif ext == ".pdf":
                return self._parse_pdf(file_bytes, filename)
            elif ext == ".txt":
                return self._parse_txt(file_bytes, filename)
            elif ext in [".png", ".jpg", ".jpeg", ".gif", ".webp"]:
                return self._parse_image(file_bytes, filename)
        except Exception as e:
            logger.error(f"解析文件失败 {filename}: {e}")
            return {
                "success": False,
                "text": "",
                "filename": filename,
                "error": f"解析失败: {str(e)}",
            }

        return {
            "success": False,
            "text": "",
            "filename": filename,
            "error": "未知错误",
        }

    def _parse_docx(self, file_bytes: bytes, filename: str) -> dict:
        """解析Word文档"""
        if not self._docx_available:
            return {
                "success": False,
                "text": "",
                "filename": filename,
                "error": "python-docx未安装，请运行: pip install python-docx",
            }

        from io import BytesIO
        try:
            doc = docx.Document(BytesIO(file_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    if any(row_text):
                        tables_text.append(" | ".join(row_text))
            full_text = "\n\n".join(paragraphs + tables_text)
            return {
                "success": True,
                "text": full_text,
                "filename": filename,
                "error": "",
            }
        except Exception as e:
            return {
                "success": False,
                "text": "",
                "filename": filename,
                "error": f"Word解析失败: {str(e)}",
            }

    def _parse_pdf(self, file_bytes: bytes, filename: str) -> dict:
        """解析PDF文档，支持OCR回退处理扫描件"""
        if not self._pdf_available:
            return {
                "success": False,
                "text": "",
                "filename": filename,
                "error": "pdfplumber未安装，请运行: pip install pdfplumber",
            }

        from io import BytesIO
        try:
            with pdfplumber.open(BytesIO(file_bytes)) as pdf:
                pages_text = []
                has_text = False
                for page in pdf.pages:
                    text = page.extract_text()
                    if text and text.strip():
                        has_text = True
                        pages_text.append(text)
                
                if has_text:
                    full_text = ""
                    for i, text in enumerate(pages_text, 1):
                        if i > 1:
                            full_text += f"\n\n--- 第{i}页 ---\n\n"
                        full_text += text
                    return {
                        "success": True,
                        "text": full_text,
                        "filename": filename,
                        "error": "",
                        "method": "pdfplumber",
                    }
                else:
                    logger.info(f"PDF {filename} 无文本内容，尝试OCR识别")
                    return self._parse_pdf_with_ocr(file_bytes, filename)
        except Exception as e:
            logger.warning(f"pdfplumber解析失败 {filename}: {e}，尝试OCR")
            return self._parse_pdf_with_ocr(file_bytes, filename)

    def _parse_pdf_with_ocr(self, file_bytes: bytes, filename: str) -> dict:
        """通过OCR解析PDF（扫描件或复杂PDF）"""
        if not self._fitz_available:
            return {
                "success": False,
                "text": "",
                "filename": filename,
                "error": "pymupdf未安装，请运行: pip install pymupdf",
            }

        if not self._llm_service:
            return {
                "success": False,
                "text": "",
                "filename": filename,
                "error": "未配置LLM服务，无法进行OCR识别",
            }

        import fitz
        from io import BytesIO

        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            all_pages_text = []
            success_count = 0
            failed_count = 0

            for page_num in range(len(doc)):
                page = doc[page_num]
                pix = page.get_pixmap(dpi=300)
                
                img_bytes = pix.tobytes("png")
                processed_img_bytes = self._preprocess_image(img_bytes)
                img_base64 = base64.b64encode(processed_img_bytes).decode("utf-8")
                
                ocr_result = self._llm_service.ocr_image(img_base64, "png")
                
                if ocr_result and ocr_result.get("success"):
                    page_text = ocr_result.get("text", "")
                    if page_text.strip():
                        all_pages_text.append(page_text)
                        success_count += 1
                    else:
                        failed_count += 1
                    logger.info(f"PDF第{page_num+1}页OCR识别成功")
                else:
                    failed_count += 1
                    logger.warning(f"PDF第{page_num+1}页OCR识别失败: {ocr_result.get('error', '未知错误') if ocr_result else '无结果'}")

            if all_pages_text:
                full_text = ""
                for i, text in enumerate(all_pages_text, 1):
                    if i > 1:
                        full_text += f"\n\n--- 第{i}页 ---\n\n"
                    full_text += text
                
                warning_msg = ""
                if failed_count > 0:
                    warning_msg = f" （{failed_count}页识别失败）"
                
                return {
                    "success": True,
                    "text": full_text,
                    "filename": filename,
                    "error": warning_msg,
                    "method": "ocr",
                    "pages": {"total": len(doc), "success": success_count, "failed": failed_count},
                }
            else:
                return {
                    "success": False,
                    "text": "",
                    "filename": filename,
                    "error": f"OCR识别失败，{len(doc)}页均无法提取文本",
                }
        except Exception as e:
            logger.error(f"OCR解析失败 {filename}: {e}", exc_info=True)
            return {
                "success": False,
                "text": "",
                "filename": filename,
                "error": f"OCR解析失败: {str(e)}",
            }

    def _preprocess_image(self, img_bytes: bytes) -> bytes:
        """图片预处理：压缩、降噪、增强对比度"""
        if not self._pil_available:
            return img_bytes
        
        try:
            from PIL import Image, ImageEnhance
            from io import BytesIO
            
            img = Image.open(BytesIO(img_bytes))
            
            if img.mode != "RGB":
                img = img.convert("RGB")
            
            max_size = 2048
            width, height = img.size
            if width > max_size or height > max_size:
                ratio = min(max_size / width, max_size / height)
                new_size = (int(width * ratio), int(height * ratio))
                img = img.resize(new_size, Image.LANCZOS)
            
            enhancer = ImageEnhance.Contrast(img)
            img = enhancer.enhance(1.2)
            
            enhancer = ImageEnhance.Sharpness(img)
            img = enhancer.enhance(1.1)
            
            buffer = BytesIO()
            img.save(buffer, format="PNG", quality=80)
            return buffer.getvalue()
        except Exception as e:
            logger.warning(f"图片预处理失败，使用原始图片: {e}")
            return img_bytes

    def _parse_txt(self, file_bytes: bytes, filename: str) -> dict:
        """解析纯文本文件"""
        try:
            text = file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            text = file_bytes.decode("gbk", errors="replace")
        return {
            "success": True,
            "text": text,
            "filename": filename,
            "error": "",
        }

    def _parse_image(self, file_bytes: bytes, filename: str) -> dict:
        """解析图片文件（通过OCR识别）"""
        if not self._pil_available:
            return {
                "success": False,
                "text": "",
                "filename": filename,
                "error": "Pillow未安装，请运行: pip install Pillow",
            }

        if not self._llm_service:
            return {
                "success": False,
                "text": "",
                "filename": filename,
                "error": "未配置LLM服务，无法进行OCR识别",
            }

        try:
            from PIL import Image
            from io import BytesIO

            img = Image.open(BytesIO(file_bytes))
            if img.mode != "RGB":
                img = img.convert("RGB")
            
            buffer = BytesIO()
            img.save(buffer, format="PNG")
            img_bytes = buffer.getvalue()
            img_base64 = base64.b64encode(img_bytes).decode("utf-8")

            ocr_result = self._llm_service.ocr_image(img_base64, "png")
            
            if ocr_result and ocr_result.get("success"):
                return {
                    "success": True,
                    "text": ocr_result.get("text", ""),
                    "filename": filename,
                    "error": "",
                    "method": "ocr",
                }
            else:
                return {
                    "success": False,
                    "text": "",
                    "filename": filename,
                    "error": f"OCR识别失败: {ocr_result.get('error', '未知错误')}",
                }
        except Exception as e:
            return {
                "success": False,
                "text": "",
                "filename": filename,
                "error": f"图片解析失败: {str(e)}",
            }

    def parse_multiple(self, files: list) -> dict:
        """
        解析多个文件并合并结果
        
        Args:
            files: 列表，每个元素包含 {"bytes": bytes, "filename": str}
        
        Returns:
            dict: {"combined_text": str, "files": list[dict], "errors": list[str]}
        """
        results = []
        errors = []
        sections = []

        for file_info in files:
            file_bytes = file_info.get("bytes", b"")
            filename = file_info.get("filename", "unknown")
            if not file_bytes:
                continue

            result = self.parse_file(file_bytes, filename)
            results.append(result)

            if result["success"]:
                sections.append(f"=== {filename} ===\n\n{result['text']}\n\n")
            else:
                errors.append(result["error"])

        combined_text = "".join(sections).strip()

        return {
            "combined_text": combined_text,
            "files": results,
            "errors": errors,
        }

    def can_parse(self, filename: str) -> bool:
        """判断是否可以解析该文件"""
        ext = os.path.splitext(filename)[1].lower()
        if ext not in SUPPORTED_EXTENSIONS:
            return False
        if ext == ".docx" and not self._docx_available:
            return False
        if ext == ".pdf" and not self._pdf_available:
            return False
        if ext in [".png", ".jpg", ".jpeg", ".gif", ".webp"] and not self._pil_available:
            return False
        return True


_thread_local = threading.local()


def get_thread_local_parser(llm_service=None) -> DocumentParser:
    """获取线程本地的DocumentParser实例（每个线程独立）"""
    if not hasattr(_thread_local, 'document_parser'):
        _thread_local.document_parser = DocumentParser(llm_service)
    elif llm_service is not None:
        _thread_local.document_parser.set_llm_service(llm_service)
    return _thread_local.document_parser


def parse_document(file_bytes: bytes, filename: str, llm_service=None) -> dict:
    """便捷函数：解析单个文件"""
    return get_thread_local_parser(llm_service).parse_file(file_bytes, filename)


def parse_documents(files: list, llm_service=None) -> dict:
    """便捷函数：解析多个文件"""
    return get_thread_local_parser(llm_service).parse_multiple(files)