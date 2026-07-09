import io
from docx import Document
from exporters.canonical import normalize
from exporters.boundary import MAX_LIST_ITEMS, MAX_TEXT_LEN
from exporters.html_exporter import generate_html
from exporters.markdown_exporter import MarkdownExporter
from exporters.ppt_spec_exporter import generate_ppt_spec
from exporters.word_exporter import WordExporter


def test_normalize_huge_list_capped():
    bs = {"objectives": [{"objective": f"o{i}", "priority": "high"} for i in range(500)]}
    r = normalize(bs)
    # 封顶 MAX_LIST_ITEMS 条真实数据 + 1 条"已省略"合成条目
    assert len(r.objectives) == MAX_LIST_ITEMS + 1
    assert any("其余" in o.objective and "已省略" in o.objective for o in r.objectives)


def test_normalize_long_text_truncated():
    bs = {"business_domain": "x" * (MAX_TEXT_LEN + 100)}
    r = normalize(bs)
    assert "已截断" in r.title


def test_normalize_none_field_becomes_placeholder():
    bs = {"roles": [{"role": "CEO", "department": None}]}
    r = normalize(bs)
    assert r.roles[0].department == "—"


def test_normalize_type_error_coerced():
    bs = {"metrics": [{"name": 12345, "formula": "x", "target": "y"}]}
    r = normalize(bs)
    assert r.metrics[0].name == "12345"


def test_normalize_control_chars_stripped():
    bs = {"business_domain": "a" + chr(0) + "b" + chr(31) + "c"}
    r = normalize(bs)
    assert chr(0) not in r.title and chr(31) not in r.title


def test_normalize_bom_stripped():
    bs = {"business_domain": chr(0xFEFF) + "项目"}
    r = normalize(bs)
    assert chr(0xFEFF) not in r.title


def test_normalize_empty_input_safe():
    r = normalize({})
    assert r.title == "业务系统分析报告"
    assert r.objectives == []


def test_html_injection_escaped():
    bs = {"business_domain": "<script>alert(1)</script>",
          "objectives": [{"objective": "<img src=x onerror=alert(1)>", "priority": "high"}]}
    r = normalize(bs)
    html = generate_html(r, {}, None)
    assert "<script>" not in html
    assert "<img" not in html
    assert "&lt;script&gt;" in html
    assert "&lt;img" in html


def _render_all(bs):
    r = normalize(bs)
    md = MarkdownExporter().export(r, None)
    html = generate_html(r, {}, None)
    ppt = generate_ppt_spec(r, None)
    doc = Document(io.BytesIO(WordExporter().export(r)))
    word = "\n".join(p.text for p in doc.paragraphs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                word += "\n" + cell.text
    return md, html, ppt, word


def _ppt_has(ppt, text):
    for s in ppt["slides"]:
        if text in s.get("title", ""):
            return True
        for it in s.get("items", []):
            if text in it:
                return True
        for row in s.get("data", []):
            if any(text in str(c) for c in row):
                return True
    return False


def test_boundary_huge_list_all_formats():
    bs = {"metrics": [{"name": f"m{i}", "formula": "x", "target": "y"} for i in range(1000)]}
    md, html, ppt, word = _render_all(bs)
    for out in (md, html, word):
        assert "其余" in out and "已省略" in out
    assert _ppt_has(ppt, "已省略")  # 指标省略项落在 PPT 表格 data 中


def test_boundary_long_text_truncated_marker():
    bs = {"business_domain": "超长" * 1500}
    md, html, ppt, word = _render_all(bs)
    for out in (md, html, word):
        assert "已截断" in out


def test_boundary_special_chars_no_crash():
    bs = {"business_domain": "</>&%" + chr(0x1F600) + chr(0x200B) + chr(0x202E) + "B",
          "objectives": [{"objective": "</>&%", "priority": "high"}]}
    md, html, ppt, word = _render_all(bs)   # 不抛即达标
    assert md and html and word


def test_boundary_html_injection_not_executable():
    bs = {"business_domain": "<script>alert(1)</script>"}
    _, html, _, _ = _render_all(bs)
    assert "<script>" not in html and "&lt;script&gt;" in html


def test_boundary_control_chars_stripped():
    bs = {"business_domain": "a" + chr(0) + "b" + chr(31) + "c"}
    md, html, ppt, word = _render_all(bs)
    for out in (md, html, word):
        assert chr(0) not in out and chr(31) not in out


def test_boundary_missing_field_safe():
    bs = {"objectives": [{"objective": "o"}]}  # 无 priority_label
    r = normalize(bs)
    assert r.objectives[0].priority_label  # 有默认标签，不崩


def test_boundary_type_error_coerced():
    bs = {"metrics": [{"name": 999, "formula": ["not", "str"], "target": "y"}]}
    md, html, ppt, word = _render_all(bs)
    assert "999" in md and "999" in html


def test_boundary_none_value_placeholder():
    bs = {"roles": [{"role": "CEO", "department": None, "level": None, "headcount": None}]}
    md, html, ppt, word = _render_all(bs)
    for out in (md, html, word):
        assert "—" in out


def test_boundary_bom_normalized():
    bs = {"business_domain": chr(0xFEFF) + "项目"}
    md, html, ppt, word = _render_all(bs)
    for out in (md, html, word):
        assert chr(0xFEFF) not in out


def test_boundary_empty_input_safe():
    md, html, ppt, word = _render_all({})
    for out in (md, html, word):
        assert out  # 有输出、不崩、段落集完整
    assert any(s["title"] == "业务目标" for s in ppt["slides"])


def test_boundary_nested_anomaly_safe():
    bs = {"risk": "not-a-list-or-dict"}  # 异常类型
    r = normalize(bs)  # 不抛即达标
    assert isinstance(r.risks, list)


def test_boundary_cross_format_section_set_intact():
    bs = {"metrics": [{"name": f"m{i}", "formula": "x", "target": "y"} for i in range(500)]}
    md, html, ppt, word = _render_all(bs)
    markers = ["业务目标", "角色定义", "业务流程", "关键指标", "风险分析", "战略建议"]
    for m in markers:
        assert m in md and m in html and m in word
    assert all(any(s["title"] == m for s in ppt["slides"]) for m in markers)
