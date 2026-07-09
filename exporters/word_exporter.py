"""Word Exporter - 生成Word文档报告"""
from typing import Dict, Any
import io
import logging

logger = logging.getLogger(__name__)


class WordExporter:
    """Word文档导出器"""

    def __init__(self):
        self._docx_available = False
        try:
            import docx
            self._docx_available = True
        except ImportError:
            logger.warning("python-docx not installed, Word export will be disabled")

    def export(self, report) -> bytes:
        """导出为Word文档。report 为 CanonicalReport。"""
        if not self._docx_available:
            from exporters.errors import ExportDependencyError
            raise ExportDependencyError("word", "python-docx", "pip install python-docx")
        from exporters.canonical import CanonicalReport, normalize
        if not isinstance(report, CanonicalReport):
            report = normalize(report)

        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.section import WD_SECTION
        from docx.oxml.ns import qn

        doc = Document()

        style = doc.styles['Normal']
        style.font.name = '微软雅黑'
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        section = doc.sections[0]
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

        title_para = doc.add_heading(report.title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if report.executive_summary:
            doc.add_paragraph(report.executive_summary, style='Intense Quote')

        doc.add_heading('一、业务目标', level=1)
        if report.objectives:
            for o in report.objectives:
                p = doc.add_paragraph(f"{o.priority_label}{o.objective}")
                if o.target:
                    p.add_run(f" - 目标: {o.target}")
        else:
            doc.add_paragraph("暂无业务目标")

        doc.add_heading('二、角色定义', level=1)
        if report.roles:
            table = doc.add_table(rows=1, cols=4)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '角色名称'
            hdr_cells[1].text = '所属部门'
            hdr_cells[2].text = '级别'
            hdr_cells[3].text = '人数'
            for r in report.roles:
                row_cells = table.add_row().cells
                row_cells[0].text = r.role
                row_cells[1].text = r.department
                row_cells[2].text = r.level
                row_cells[3].text = str(r.headcount)
        else:
            doc.add_paragraph("暂无角色定义")

        doc.add_heading('三、业务流程', level=1)
        if report.workflow:
            for s in report.workflow:
                p = doc.add_paragraph(f"{s.step}. {s.name}", style='List Number')
                if s.action:
                    p.add_run(f"\n   动作: {s.action}")
                if s.role:
                    p.add_run(f"\n   负责角色: {s.role}")
        else:
            doc.add_paragraph("暂无业务流程")

        doc.add_heading('四、关键指标', level=1)
        if report.metrics:
            for m in report.metrics:
                line = m.name
                if m.formula:
                    line += f"（公式: {m.formula}）"
                if m.target:
                    line += f" 目标: {m.target}"
                doc.add_paragraph(line)
        else:
            doc.add_paragraph("暂无关键指标")

        doc.add_heading('五、风险分析', level=1)
        if report.risks:
            for rk in report.risks:
                p = doc.add_paragraph(f"{rk.severity_label}: {rk.risk}")
                if rk.mitigation:
                    p.add_run(f"\n   应对措施: {rk.mitigation}")
                if rk.impact:
                    p.add_run(f"\n   影响: {rk.impact}")
        else:
            doc.add_paragraph("暂无风险分析")

        doc.add_heading('六、战略建议', level=1)
        for rec in report.strategy.recommendations:
            doc.add_paragraph(rec, style='List Bullet')
        for g in report.strategy.growth_opportunities:
            doc.add_paragraph(f"{g['opportunity']}: {g['potential']}")
        for step in report.strategy.roadmap:
            doc.add_paragraph(step, style='List Bullet')
        if not (report.strategy.recommendations or report.strategy.growth_opportunities or report.strategy.roadmap):
            doc.add_paragraph("暂无战略建议")

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)

        return output.read()
